import os
import re
import json
from typing import List
from flask import Flask, render_template, request, jsonify
from docx import Document
import vertexai
from vertexai.generative_models import GenerativeModel
from database import SessionLocal, engine
from models import Base, Interaction

app = Flask(__name__, template_folder="templates", static_folder="static")

SERVICE_ACCOUNT_PATH = "intern-service-account.json"
os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = SERVICE_ACCOUNT_PATH
PROJECT_ID = "ela-cert-test"
LOCATION = "us-central1"
MODEL_NAME = "gemini-1.5-pro-002"
DOWNLOAD_DIR = "./downloads/email_bot/Email Writing"

vertexai.init(project=PROJECT_ID, location=LOCATION)
Base.metadata.create_all(bind=engine)

def extract_text_from_docx(file_path: str) -> str:
    doc = Document(file_path)
    text_blocks: List[str] = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_blocks.append(cell.text.strip())
    return "\n".join(text_blocks)

def split_into_examples(text: str) -> List[str]:
    pattern = r"(?i)(Example\s*\d+\s*:?)"
    matches = list(re.finditer(pattern, text))
    out: List[str] = []
    for i, m in enumerate(matches):
        start = m.start()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        chunk = text[start:end].strip()
        if chunk:
            out.append(chunk)
    return out

def generate_new_question(content: str, scenario: str, cefr_level: str, existing_questions: List[str], recent_questions: List[str]) -> str | None:
    model = GenerativeModel(MODEL_NAME)
    prompt = f"""
You are an English exam question creator.

Scenario: {scenario}
CEFR Level: {cefr_level}

Existing Questions:
{existing_questions}

Content for reference:
{content}

Task:
Generate one unique scenario-based email-writing question based on the given content and CEFR level.
The question must NOT duplicate any from Existing Questions or Recent Questions.
Keep it concise and engaging.
"""
    resp = model.generate_content(prompt)
    new_q = getattr(resp, "text", "").strip()
    if new_q and new_q not in existing_questions and new_q not in recent_questions:
        recent_questions.append(new_q)
        if len(recent_questions) > 5:
            recent_questions.pop(0)
        return new_q
    return None

def evaluate_email_model(email_content: str, scenario: str, scenario_question: str, cefr_level: str) -> dict:
    model = GenerativeModel(MODEL_NAME)
    prompt = f"""
You are an English writing evaluator.
Scenario: {scenario}
Scenario Question: {scenario_question}
CEFR Level: {cefr_level}
Email content to evaluate:
{email_content}

Evaluate the email based on:
- Greeting
- Body clarity and relevance
- Sign-off
- Grammar and vocabulary

Give a rating from 1 to 5 and detailed feedback.
"""
    resp = model.generate_content(prompt)
    feedback_text = getattr(resp, "text", "").strip()
    format_eval = {
        "greeting": bool(re.search(r"dear|hello|hi", email_content, re.IGNORECASE)),
        "body": len(email_content.split()) > 20,
        "sign_off": bool(re.search(r"regards|sincerely|thank you", email_content, re.IGNORECASE)),
    }
    rating = min(5, max(1, feedback_text.count(".") % 6))
    return {"feedback": feedback_text, "rating": rating, "format_evaluation": format_eval}

@app.get("/")
def home():
    return render_template("index.html")

@app.get("/health")
def health():
    return {"ok": True}

@app.get("/api/files")
def list_files():
    try:
        files = [f for f in os.listdir(DOWNLOAD_DIR) if f.lower().endswith(".docx")]
        files.sort()
        return jsonify(files)
    except Exception as e:
        return jsonify({"detail": str(e)}), 500

@app.post("/generate_questions")
def generate_questions_endpoint():
    data = request.get_json(force=True) or {}
    file_path = data.get("file_path", "")
    scenario = data.get("scenario", "")
    cefr_level = data.get("cefr_level", "")
    existing_questions = data.get("existing_questions", []) or []
    if not file_path or not scenario or not cefr_level:
        return jsonify({"detail": "file_path, scenario, cefr_level are required"}), 400
    abs_path = os.path.join(DOWNLOAD_DIR, file_path)
    if not os.path.exists(abs_path):
        return jsonify({"detail": f"File not found: {abs_path}"}), 404
    full_content = extract_text_from_docx(abs_path)
    examples = split_into_examples(full_content)
    if not examples:
        return jsonify({"detail": "No examples found in file"}), 400
    recent_questions: List[str] = []
    new_questions: List[str] = []
    for ex in examples:
        q = generate_new_question(ex, scenario, cefr_level, existing_questions, recent_questions)
        if q:
            new_questions.append(q)
    resp = {"new_questions": new_questions}
    sess = SessionLocal()
    try:
        row = Interaction(
            kind="questions",
            scenario=scenario,
            cefr_level=cefr_level,
            request_json=json.dumps(data, ensure_ascii=False),
            response_json=json.dumps(resp, ensure_ascii=False),
        )
        sess.add(row)
        sess.commit()
    finally:
        sess.close()
    return jsonify(resp)

@app.post("/evaluate_email")
def evaluate_email_endpoint():
    data = request.get_json(force=True) or {}
    email_content = data.get("email_content", "")
    scenario = data.get("scenario", "")
    scenario_question = data.get("scenario_question", "")
    cefr_level = data.get("cefr_level", "")
    if not email_content or not scenario or not scenario_question or not cefr_level:
        return jsonify({"detail": "email_content, scenario, scenario_question, cefr_level are required"}), 400
    result = evaluate_email_model(email_content, scenario, scenario_question, cefr_level)
    resp = result
    sess = SessionLocal()
    try:
        row = Interaction(
            kind="email",
            scenario=scenario,
            cefr_level=cefr_level,
            request_json=json.dumps(data, ensure_ascii=False),
            response_json=json.dumps(resp, ensure_ascii=False),
        )
        sess.add(row)
        sess.commit()
    finally:
        sess.close()
    return jsonify(resp)

@app.get("/api/logs")
def api_logs():
    sess = SessionLocal()
    try:
        rows = (
            sess.query(Interaction)
            .order_by(Interaction.id.desc())
            .limit(20)
            .all()
        )
        out = []
        for r in rows:
            out.append({
                "id": r.id,
                "kind": r.kind,
                "scenario": r.scenario,
                "cefr_level": r.cefr_level,
                "request_json": r.request_json,
                "response_json": r.response_json,
                "created_at": r.created_at.isoformat() if r.created_at else None
            })
        return jsonify(out)
    finally:
        sess.close()

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=8000, debug=True)
